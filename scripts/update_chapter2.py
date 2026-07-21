"""
Python Script to Clean, Condense, and Inject Chapter 2 Content into Word Document

This script:
1. Reads source text from a file
2. Condenses content according to strict layout constraints
3. Creates a backup of the target Word document
4. Replaces Chapter 2 content in the target document
5. Applies proper formatting and red placeholders

Author: Automation Engineer
Date: 2025
"""

import re
import shutil
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# ============== CONFIGURATION ==============
SOURCE_FILE = r"C:\Users\Youcode\Downloads\Rapport_Stage_Devosoft\Chapitre2_Etat_De_L_Art_Complet.txt"
TARGET_FILE = r"C:\Users\Youcode\Downloads\Rapport_Stage_Devosoft\Rapport_Final_FactureX_New.docx"

# ============== CONDENSATION RULES ==============
MAX_LINES_GENERAL = 7  # Maximum lines for general technology descriptions
MAX_LINES_UML = 6      # Maximum lines for UML section (CRITICAL)


class Chapter2Updater:
    """Handles the reading, condensation, and injection of Chapter 2 content."""

    def __init__(self, source_file, target_file):
        self.source_file = Path(source_file)
        self.target_file = Path(target_file)
        self.content = None
        self.condensed_content = []

    def read_source(self):
        """Read the source text file."""
        print(f"[READ] Reading source file: {self.source_file}")
        with open(self.source_file, 'r', encoding='utf-8') as f:
            self.content = f.readlines()
        print(f"[OK] Read {len(self.content)} lines from source file")
        return self.content

    def create_backup(self):
        """Create a backup of the target Word document."""
        # Create backup filename by inserting _BACKUP before the extension
        stem = self.target_file.stem
        backup_path = self.target_file.parent / f"{stem}_BACKUP.docx"
        print(f"[BACKUP] Creating backup: {backup_path}")
        shutil.copy2(self.target_file, backup_path)
        print(f"[OK] Backup created successfully")
        return backup_path

    def is_section_header(self, line):
        """Check if a line is a section header."""
        line = line.strip()
        if not line:
            return False
        # Check for chapter headers (2.x, 2.x.x)
        if re.match(r'^2\.?\d+(\.\d+)?\s+[A-Z]', line):
            return True
        # Check for separator lines
        if line.startswith('═') or line.startswith('='):
            return True
        return False

    def is_figure_placeholder(self, line):
        """Check if a line is a figure placeholder."""
        return bool(re.search(r'\[🔴 REMARQUE', line))

    def condense_text_section(self, lines, max_lines=MAX_LINES_GENERAL):
        """
        Condense a text section to max_lines by removing redundant content
        while preserving the core information.
        """
        if len(lines) <= max_lines:
            return lines

        # Keep the first paragraph (intro)
        result = []
        first_para = []
        for line in lines:
            first_para.append(line)
            if line.strip() == '' and len(first_para) > 1:
                break

        # Keep figure placeholders
        placeholders = [l for l in lines if self.is_figure_placeholder(l)]

        # Get key sentences from middle content
        middle_content = lines[len(first_para):-len(placeholders) if placeholders else None]
        if middle_content:
            # Take sentences that start with capital letters (topic sentences)
            key_sentences = []
            for line in middle_content:
                line = line.strip()
                if line and (line[0].isupper() or line.startswith('-') or line.startswith('•')):
                    key_sentences.append(line)
                    if len(key_sentences) >= max_lines - 2:
                        break

            result = first_para[:2] + key_sentences[:max_lines-2]
        else:
            result = first_para[:max_lines]

        # Add back placeholders
        result.extend(placeholders)

        return result

    def condense_uml_section(self, lines):
        """
        CRITICAL: Condense UML section to ONLY high-level definition.
        NO listing of diagram types, NO detailed explanations.
        """
        print("[CONDENSE] Applying CRITICAL UML condensation rules...")

        # Find the UML section
        uml_start = -1
        uml_end = -1

        for i, line in enumerate(lines):
            if 'UML' in line and ('2.7' in line or 'langage' in line.lower() or 'modélisation' in line.lower()):
                uml_start = i
                break

        if uml_start == -1:
            return lines

        # Find end of UML section (next section or figure placeholder)
        for i in range(uml_start + 1, len(lines)):
            if self.is_section_header(lines[i]) or ('Figure' in lines[i] and '═' in lines[i-1] if i > 0 else False):
                uml_end = i
                break

        if uml_end == -1:
            uml_end = len(lines)

        # Create condensed UML section (max 6 lines)
        condensed_uml = [
            lines[uml_start],  # Header
            "\n",
            "UML (Unified Modeling Language) est un langage de modélisation graphique standardisé\n",
            "par l'OMG pour la spécification, la visualisation et la documentation des systèmes\n",
            "logiciels. Il propose 14 types de diagrammes répartis en diagrammes structurels\n",
            "(statiques) et diagrammes comportementaux (dynamiques), facilitant ainsi la\n",
            "communication entre les acteurs d'un projet et l'analyse approfondie du système.\n",
            "\n",
        ]

        # Add figure placeholder if exists
        for i in range(uml_start, uml_end):
            if self.is_figure_placeholder(lines[i]):
                condensed_uml.append(lines[i])
                break

        # Replace UML section with condensed version
        result = lines[:uml_start] + condensed_uml + lines[uml_end:]
        return result

    def process_content(self):
        """
        Process and condense the source content according to rules.
        """
        print("[PROCESS] Processing content with condensation rules...")

        current_section = []
        processed = []
        in_uml_section = False

        for line in self.content:
            # Check if this is a new section
            if self.is_section_header(line) and current_section:
                # Process previous section
                if in_uml_section:
                    # Will handle UML separately
                    pass
                else:
                    condensed = self.condense_text_section(current_section)
                    processed.extend(condensed)
                current_section = []
                in_uml_section = False

            current_section.append(line)

            # Check for UML section
            if 'UML' in line and '2.7' in line:
                in_uml_section = True

        # Process last section
        if current_section:
            condensed = self.condense_text_section(current_section)
            processed.extend(condensed)

        # Apply UML condensation as a separate pass
        self.condensed_content = self.condense_uml_section(processed)

        print(f"[OK] Content processed: {len(self.content)} -> {len(self.condensed_content)} lines")
        return self.condensed_content

    def get_heading_level(self, text):
        """Determine heading level from text."""
        text = text.strip()
        if re.match(r'^2\.\d+\s+[A-Z]', text):
            return 2  # Heading 2 for 2.x sections
        elif re.match(r'^2\.\d+\.\d+\s+[A-Z]', text):
            return 3  # Heading 3 for 2.x.x sections
        elif re.match(r'^2\.?\s*[A-Z]|Chapitre 2', text, re.IGNORECASE):
            return 1  # Heading 1 for main chapter title
        return None

    def inject_into_document(self):
        """
        Inject the condensed content into the target Word document,
        replacing the existing Chapter 2.
        """
        print(f"[OPEN] Opening target document: {self.target_file}")
        doc = Document(str(self.target_file))

        # Find Chapter 2 location
        chap2_start = None
        chap3_start = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if 'Chapitre 2' in text or ('État de l\'Art' in text and 'Chapitre' in doc.paragraphs[i-1].text if i > 0 else False):
                chap2_start = i
            elif 'Chapitre 3' in text:
                chap3_start = i
                break

        if chap2_start is None:
            raise ValueError("Chapter 2 not found in target document!")

        print(f"[FOUND] Chapter 2 at paragraph {chap2_start}")
        print(f"[FOUND] Chapter 3 starts at paragraph {chap3_start}")

        # Remove old Chapter 2 content
        # We need to remove from chap2_start to chap3_start
        # But python-docx doesn't support deletion well, so we'll rebuild

        # Get document elements before Chapter 2
        new_doc_content = []

        # Keep everything before Chapter 2
        for i in range(chap2_start):
            new_doc_content.append(('keep', i))

        # Add marker for Chapter 2 replacement
        new_doc_content.append(('replace', chap2_start))

        # Keep everything from Chapter 3 onwards
        if chap3_start:
            for i in range(chap3_start, len(doc.paragraphs)):
                new_doc_content.append(('keep', i))

        # Create new document
        new_doc = Document()

        # Copy styles from original
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:
                try:
                    new_doc.styles.add_style(style.name, style.type)
                except:
                    pass

        # Process content
        is_chapter2_replacement = False

        # First, add content before Chapter 2
        for action, idx in new_doc_content:
            if action == 'keep':
                para = doc.paragraphs[idx]
                new_para = new_doc.add_paragraph(para.text, style=para.style)
                # Copy formatting
                for run in para.runs:
                    new_run = new_para.runs[-1] if new_para.runs else new_para.add_run()
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline
                    if run.font.color.rgb:
                        new_run.font.color.rgb = run.font.color.rgb
            elif action == 'replace' and not is_chapter2_replacement:
                is_chapter2_replacement = True
                # Add condensed Chapter 2 content
                self._add_condensed_content(new_doc)
                # Skip to next iteration
                continue

        # Save the modified document
        print(f"[SAVE] Saving modified document to: {self.target_file}")
        new_doc.save(str(self.target_file))
        print(f"[OK] Document updated successfully!")

    def _add_condensed_content(self, doc):
        """Add the condensed content to the document with proper formatting."""
        current_text = ""

        for line in self.condensed_content:
            line = line.rstrip()

            # Handle section separators
            if line.startswith('═') or line.startswith('='):
                if current_text.strip():
                    self._add_paragraph_with_style(doc, current_text.strip())
                    current_text = ""
                continue

            # Handle figure placeholders
            if self.is_figure_placeholder(line):
                if current_text.strip():
                    self._add_paragraph_with_style(doc, current_text.strip())
                    current_text = ""
                # Add placeholder in red
                para = doc.add_paragraph(line.strip())
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.font.bold = True
                continue

            # Build current text
            current_text += line + "\n"

            # Check if we should end the paragraph
            if line.strip() == "" and current_text.strip():
                self._add_paragraph_with_style(doc, current_text.strip())
                current_text = ""

        # Add remaining content
        if current_text.strip():
            self._add_paragraph_with_style(doc, current_text.strip())

    def _add_paragraph_with_style(self, doc, text):
        """Add a paragraph with appropriate styling based on content."""
        # Check if this is a heading
        heading_level = self.get_heading_level(text)

        if heading_level == 1:
            para = doc.add_heading(text, level=1)
        elif heading_level == 2:
            para = doc.add_heading(text, level=2)
        elif heading_level == 3:
            para = doc.add_heading(text, level=3)
        elif text.startswith('-') or text.startswith('•'):
            # Bullet point
            para = doc.add_paragraph(text[1:].strip(), style='List Bullet')
        elif text.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            # Numbered list
            para = doc.add_paragraph(text[3:].strip(), style='List Number')
        else:
            # Regular paragraph
            para = doc.add_paragraph(text)

        # Check for red placeholders within the text
        if '[🔴 REMARQUE' in text:
            for run in para.runs:
                if '[🔴 REMARQUE' in run.text:
                    run.font.color.rgb = RGBColor(255, 0, 0)
                    run.font.bold = True

    def run(self):
        """Execute the full update process."""
        print("=" * 60)
        print("CHAPTER 2 UPDATE SCRIPT")
        print("=" * 60)

        # Step 1: Read source
        self.read_source()

        # Step 2: Process and condense
        self.process_content()

        # Step 3: Create backup
        self.create_backup()

        # Step 4: Inject into document
        self.inject_into_document()

        print("=" * 60)
        print("CHAPTER 2 UPDATE COMPLETED SUCCESSFULLY!")
        print("=" * 60)


def main():
    """Main entry point."""
    updater = Chapter2Updater(SOURCE_FILE, TARGET_FILE)
    updater.run()


if __name__ == "__main__":
    main()
